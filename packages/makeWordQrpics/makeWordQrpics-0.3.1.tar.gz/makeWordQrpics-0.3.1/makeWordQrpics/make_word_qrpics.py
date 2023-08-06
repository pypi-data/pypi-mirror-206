'''
底稿目录制作隔页
'''
# -*- coding: utf-8 -*-
import re
import traceback
from PyQt5 import QtCore
import docx
from docx.oxml.ns import qn
from docx.shared import Inches
import sys
from PyQt5 import QtWidgets
try:
    from Ui.TableToWordUi import Ui_Form
    from Function.optionDb import getOption, InsertDb
except:
    from .Ui.TableToWordUi import Ui_Form
    from .Function.optionDb import getOption, InsertDb
import os
from PyQt5.QtWidgets import *
from docx.oxml.shared import OxmlElement
from PyQt5.QtCore import *

CM = 360000
from clazz import demo

class MyPyQT_Form(QtWidgets.QWidget,Ui_Form, demo.TableObject):
    #警告信号量
    WarningSignal = pyqtSignal(str, str)
    PlainTextSignal = pyqtSignal(str)
    ReloadSignal = pyqtSignal()
    ProBarSignal = pyqtSignal(float)
    def __init__(self):
        super(MyPyQT_Form,self).__init__()
        self.setupUi(self)
        self.ReloadFun()
        self.WarningSignal.connect(self.showmsg)
        self.StartBtn.clicked.connect(self.StartFun)
        self.UiInit()
        self.PlainTextSignal.connect(self.SetPlaintext)
        self.ReloadSignal.connect(self.ReloadBtn),
        self.ProBarSignal.connect(self.ProcessBarFuntion)
        self.ColorList = ['0', '1', '2', '3', '4','5', '6', '7', '8', '9', "A", "B", "C", "D", "E", "F"]
        self.WordPathBtn.clicked.connect(self.file_Path)
        self.UseList = []
        self.FunctionName = "底稿目录制作隔页"
        self.describetion = '''
=================================================
这是底稿目录制作隔页的功能
=================================================
        '''

    #获取文件路径
    def file_Path(self):
        dirPath = QtWidgets.QFileDialog.getExistingDirectory(self, "浏览", "./")
        if dirPath:
            self.WordPathEdit.setText(dirPath)


    #界面进度条函数
    def ProcessBarFuntion(self, num):
        if num < 1.1:
            self.progressBar.setValue(int(num * 100))
        else:
            self.progressBar.setValue(100)
            self.progressBar.setTextVisible(False)
            self.FinishLabel.setVisible(True)

    def SetPlaintext(self, word):
        self.plainTextEdit.setPlainText(word)

    def UiInit(self):
        ui_list = getOption("WordQrpics")
        # 判断类型
        if isinstance(ui_list, list):
            if len(ui_list) == 1:
                ui_tuple = ui_list[0]
                self.WordPathEdit.setText(ui_tuple[1])
                self.colorEdit.setText(ui_tuple[2])
                self.levelsEdit.setText(ui_tuple[3])
        elif isinstance(ui_list, str):
            pass
        else:
            self.plainTextEdit.setPlainText(str(ui_list))

    #重置所有信号
    def ReloadFun(self):
        self.WordPathTip.setVisible(False)
        self.progressBar.setValue(0)
        self.progressBar.setTextVisible(True)
        self.FinishLabel.setVisible(False)
        self.colorTip.setVisible(False)

    #提示警告信号的函数
    def showmsg(self,t,msg):
        if(t == "warning"):
            QMessageBox.warning(self,"Warining",msg,QMessageBox.Ok)
        if(t == "info"):
            QMessageBox.information(self,"info",msg,QMessageBox.Yes,QMessageBox.Yes)

    #重置开始按钮
    def ReloadBtn(self):
        self.StartBtn.setText("开始执行！")
        self.StartBtn.setEnabled(True)

    def StartRunBtn(self):
        self.StartBtn.setText("正在执行！")
        self.StartBtn.setDisabled(True)
        self.FinishLabel.setVisible(False)
        self.progressBar.setValue(0)
        self.progressBar.setTextVisible(True)
        self.WordPathTip.setVisible(False)
        self.FinishLabel.setVisible(False)
        self.colorTip.setVisible(False)

    #开始执行的功能
    def StartFun(self):
        print(self.UseList)
        #开始运行的按钮
        checkTuple = ("WordQrpics", self.WordPathEdit.text(),  self.colorEdit.text(), self.levelsEdit.text())
        g = InsertDb(checkTuple)
        self.StartRunBtn()
        # 隔页word保存位置
        self.WordSavePath = self.WordPathEdit.text()
        if len(self.WordSavePath) == 0:
            print("请输入隔页保存位置！")
            self.ReloadBtn()
            self.WordPathTip.setText('请输入隔页保存位置！')
            self.WordPathTip.setVisible(True)
            return

        if not os.path.exists(self.WordSavePath):
            print("隔页保存位置不存在")
            self.ReloadBtn()
            self.WordPathTip.setText("隔页保存位置不存在！")
            self.WordPathTip.setVisible(True)
            return

        #隔页的背景颜色
        self.BackColor = self.colorEdit.text()
        #9位rgb颜色是否已经匹配
        rgbColor = False
        #匹配rgb颜色的正则
        rgbpattern = re.compile("(2[0-5]{2}|[0-1][0-9]{2})(2[0-5]{2}|[0-1][0-9]{2})(2[0-5]{2}|[0-1][0-9]{2})")
        if re.match(rgbpattern, self.BackColor) and len(self.BackColor) == 9:
            print("颜色正确")
            RgbNum = self.BackColor
            rgbColor = self.ColorList[int(RgbNum[:3]) // 16] + self.ColorList[int(RgbNum[:3]) % 16] + self.ColorList[int(RgbNum[3:6]) // 16] + self.ColorList[int(RgbNum[3:6]) % 16]+\
                       self.ColorList[int(RgbNum[6:]) // 16] + self.ColorList[int(RgbNum[6:]) % 16]
        print("九位数字颜色码匹配到的颜色")
        print(rgbColor)
        #匹配6位颜色的正则
        ColorPattern = re.compile("[0-9A-Fa-f]{6}")
        if not rgbColor:
            if re.match(ColorPattern, self.BackColor.upper()) and len(self.BackColor) == 6:
                print("颜色正确")
                rgbColor = self.BackColor.upper()
            else:
                print("6位颜色编码")
                self.colorTip.setVisible(True)
                self.ReloadBtn()
                return
        if len(self.levelsEdit.text()) == 0:
            self.制作隔页层级 = 2
        else:
            self.制作隔页层级 = int(self.levelsEdit.text()) - 1
        self.thread1 = RunThread(self, ws,rgbColor,self.WordSavePath)
        self.thread1.start()

#执行的主线程
class RunThread(QThread):
    def __init__(self, communication, ws,rgbColor, WordSavePath):
        super(RunThread, self).__init__()
        self.communication =  communication
        self.WordSavaPath = WordSavePath
        # print('b')
        # wb = load_workbook(r'C:\Users\RD\Desktop\中科软\变动隔页.xlsx')  # 不用改动
        # ws = wb['跨册']  # 建目录的底稿目录表的Sheet名
        # layers = 3  # 要建立的文件夹层次，最末一级目录有几个杠就写几，如果要建到最末一级，就填写一个比较大的数字比如9999
        # start_tree = 1  # 底稿目录开始建文件夹的行，注意选中范围必须是单个数字打头
        # end_tree = 15  # 底稿目录EndLine
        # x = docx.Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        self.ws = ws
        self.rgbColor = rgbColor

    def run(self):
        layers = self.communication.制作隔页层级# 要建立的文件夹层次，最末一级目录有几个杠就写几，如果要建到最末一级，就填写一个比较大的数字比如9999
        try:
            root_path = self.WordSavaPath + '\\' + 'geye.docx'  # 输出文件目录的路径
            try:
                # print("要记得修改")
                # x = docx.Document(docx=r'F:\0工作\00old\pdf_spliter_master-master\officeTemp\default.docx')
                x = docx.Document(docx=os.getcwd() + r'\officeTemp\default.docx')#
            except:
                self.communication.WarningSignal.emit("warning", "程序目录下，default.docx文档读取失败")
            if len(self.communication.UseList) == 0:
                self.communication.WarningSignal.emit("warining", "请上传底稿目录！")
                return
            tbs = []
            for i in self.communication.UseList:
                if i[0] != "":
                    tbs.append([i[0], i[1] + i[2],1])
            tbsCopy = tbs.copy()
            i = 0
            for y in tbs:
                self.communication.ProBarSignal.emit((tbsCopy.index(y) + 1) / len(tbs))
                print((tbsCopy.index(y) + 1) / len(tbs))
                # tbsCopy[tbsCopy.index(y)] = "-----------------------------------"
                if y[0]:
                    n = y[0].count('-')
                try:
                    demoxxxxx=(int(y[2]) + 1)
                except:
                    self.communication.PlainTextSignal.emit("D列读取到非整数内容，"+str(y[0])+'  '+str(y[1])+'  '+str(y[2])+"请复查excel表")
                if y[1] and y[0] and n <= layers:
                    for m in range(1, int(y[2]) + 1):
                        i += 1
                        if m == 1:
                            result = y[0] + "  " + y[1]
                        else:
                            result = y[0] + "  " + y[1] + '（续' + str(m - 1) + ''

                        x.add_paragraph(result)
                        x.add_page_break()

            x.styles['Normal'].font.name = u'黑体'
            x.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            x_ps = x.paragraphs
            for i in range(0, len(x_ps)):
                x_ps[i].paragraph_format.alignment = 1
                x_p_runs = x_ps[i].runs
                for j in range(0, len(x_p_runs)):
                    x_p_runs[j].font.size = docx.shared.Pt(16)
                    if i == len(x_ps) - 1 and j == len(x_p_runs) - 1:
                        x_p_runs[j].clear()

            x_sts = x.sections
            for i in range(0, len(x_sts)):
                x_sts[i].top_margin = docx.shared.Cm(10)
                x_sts[i].left_margin = 3 * CM
                x_sts[i].right_margin = 3 * CM
                x_sts[i].bottom_margin = 3 * CM
            shd = OxmlElement('w:background')
            #设置背景颜色
            shd.set(qn('w:color'), '{}'.format(self.rgbColor))
            x.element.insert(0, shd)

            shd1 = OxmlElement('w:displayBackgroundShape')
            x.settings.element.insert(0, shd1)
            x.save(root_path)
            success = True
        except Exception as e:
            success = False
            traceback.print_exc()
            self.communication.ProBarSignal.emit(2)
            if  'Permission denied' in str(e):
                success = False
                self.communication.PlainTextSignal.emit("请关闭{}文件！".format(root_path))
                self.communication.ReloadSignal.emit()
            else:
                self.communication.PlainTextSignal.emit("warning", "程序报错，停止运行")
                self.communication.ReloadSignal.emit()
        if success:
            self.communication.PlainTextSignal.emit('word版隔页，详见geye.docx.')
            self.communication.ProBarSignal.emit(2)
        self.communication.ReloadSignal.emit()
    # self.textEdit.setText("你点击了按钮")


if __name__ == '__main__':
    QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling)
    app = QtWidgets.QApplication(sys.argv)
    my_pyqt_form = MyPyQT_Form()
    my_pyqt_form.show()
    sys.exit(app.exec_())
