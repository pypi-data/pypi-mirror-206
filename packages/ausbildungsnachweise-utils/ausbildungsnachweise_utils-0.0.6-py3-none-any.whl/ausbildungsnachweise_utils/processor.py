from typing import Dict

import fitz
import shutil
import subprocess
import xml.etree.ElementTree as ET

from pathlib import Path
from docx import Document


FILE_DIR = Path(__file__).parent

LIBREOFFICE_BINARY = "libreoffice"

DEFAULT_SIGNATURE_FILE_PATH = FILE_DIR / "../resources/default_signature.png"
DEFAULT_TEMPLATE_FILE_PATH = FILE_DIR / "../resources/template_ausbildungsnachweis.docx"


def sign_pdf(input_file_path: Path, output_file_path: Path, signature_img_file_path: Path = DEFAULT_SIGNATURE_FILE_PATH) -> None:
    doc = fitz.open(input_file_path)

    for page in doc:
        # Find
        text = "Ausbildungsbeauftragte"
        text_instances = page.search_for(text)

        for inst in text_instances:
            # print(inst)

            xm = (inst.x0 / 72) * 2.54
            ym = (inst.y0 / 72) * 2.54

            ym -= 200

            img_rect = fitz.Rect(inst.x0 - 30, inst.y0 - 50, inst.x1 + 30, inst.y1)

            # print(img_rect)
            page.insert_image(img_rect, filename=signature_img_file_path)

    doc.save(output_file_path, garbage=4, deflate=True, clean=True)


def read_input(input_file_path: Path) -> Dict[str, str]:
    element_tree = ET.parse(input_file_path)
    root_node = element_tree.getroot()

    input: Dict[str, str] = {
        "vorname": root_node.find("vorname").text,
        "name": root_node.find("name").text,
        "nr": root_node.find("nr").text,
        "lehrjahr": root_node.find("lehrjahr").text,
        "ausbildungswoche_von": root_node.find("ausbildungswoche_von").text,
        "ausbildungswoche_bis": root_node.find("ausbildungswoche_bis").text,
        "abteilung": root_node.find("abteilung").text,
        "betriebliche_taetigkeit": root_node.find("betriebliche_taetigkeit").text,
        "betrieblicher_unterricht": root_node.find("betrieblicher_unterricht").text,
        "berufsschule": root_node.find("berufsschule").text
    }
    return input


def fill_template(input_file_path: Path, output_file_path: Path, template_file_path: Path = DEFAULT_TEMPLATE_FILE_PATH) -> None:
    shutil.copy(template_file_path, output_file_path)
    document = Document(output_file_path)

    input = read_input(input_file_path)

    document.tables[0].cell(1, 2).text = f"Ausbildungsnachweis Nr.: {input['nr']}"
    document.tables[0].cell(2, 2).text = f"{input['name']}, {input['vorname']}"
    document.tables[0].cell(0, 5).text = input["lehrjahr"]
    document.tables[0].cell(1, 5).text = input["ausbildungswoche_von"]
    document.tables[0].cell(2, 5).text = input["ausbildungswoche_bis"]
    document.tables[0].cell(3, 5).text = input["abteilung"]
    document.tables[0].cell(5, 5).text = input["betriebliche_taetigkeit"]
    document.tables[0].cell(7, 5).text = input["betrieblicher_unterricht"]
    document.tables[0].cell(9, 5).text = input["berufsschule"]

    document.save(output_file_path)


def convert_docx_to_pdf(input_file_path: Path, output_dir_path: Path) -> None:
    subprocess.run([LIBREOFFICE_BINARY, "--headless", "--convert-to", "pdf", "--outdir", output_dir_path, input_file_path])


